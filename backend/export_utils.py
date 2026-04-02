# backend/export_utils.py — PDF & DOCX generation
import io
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor


def generate_pdf(nl_query: str, sql_query: str, table_name: str,
                 intent: str, rows: list, columns: list) -> bytes:
    """Returns PDF as bytes."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle('ClamkaTitle', parent=styles['Heading1'],
                                 textColor=colors.HexColor('#6366f1'))
    mono_style = ParagraphStyle('ClamkaMono', parent=styles['Code'],
                                fontName='Courier', fontSize=9,
                                backColor=colors.HexColor('#f0f0f0'),
                                textColor=colors.HexColor('#333333'))

    story.append(Paragraph("ClamkaBot — Generated Query", title_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        f"Table: {table_name} | Intent: {intent} | "
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        styles['Normal']
    ))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Natural Language Query", styles['Heading2']))
    story.append(Paragraph(nl_query, styles['Normal']))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Generated SQL", styles['Heading2']))
    story.append(Preformatted(sql_query, mono_style))

    if rows:
        story.append(Spacer(1, 8))
        story.append(Paragraph(f"Results Preview ({len(rows)} rows)", styles['Heading2']))
        preview = " | ".join(columns) + "\n"
        for r in rows[:10]:  # first 10 rows in PDF
            preview += " | ".join(str(v) for v in r.values()) + "\n"
        story.append(Preformatted(preview, mono_style))

    doc.build(story)
    return buf.getvalue()


def generate_docx(nl_query: str, sql_query: str, table_name: str,
                  intent: str, rows: list, columns: list) -> bytes:
    """Returns DOCX as bytes."""
    doc = Document()
    doc.add_heading("ClamkaBot — Generated Query", 0)
    doc.add_paragraph(
        f"Table: {table_name}  |  Intent: {intent}  |  "
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    doc.add_heading("Natural Language Query", 2)
    doc.add_paragraph(nl_query)
    doc.add_heading("Generated SQL", 2)
    p = doc.add_paragraph()
    run = p.add_run(sql_query)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)

    if rows:
        doc.add_heading(f"Results Preview ({len(rows)} rows)", 2)
        table = doc.add_table(rows=1 + len(rows[:10]), cols=len(columns))
        for i, col in enumerate(columns):
            table.cell(0, i).text = col
        for ri, row in enumerate(rows[:10]):
            for ci, val in enumerate(row.values()):
                table.cell(ri + 1, ci).text = str(val)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
